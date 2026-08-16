"""
thinking_log.py

Writes down what happened inside the agent, in a Word file anybody can open.

The decision records saved by decision_log.py hold everything, but they are
stored the way a computer likes to read them. This writes the same story the
way a person likes to read it, and adds each new run underneath the last one,
so the file grows into a history you can scroll back through.

Run the agent and this is written automatically. Open it with:
    data/how_the_agent_thinks.docx
"""

import os
from datetime import datetime

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

FILE_PATH = os.path.join("data", "how_the_agent_thinks.docx")
MARKDOWN_PATH = os.path.join("data", "how_the_agent_thinks.md")

# colours, kept to violet and grey so it reads calmly
DEEP = RGBColor(0x4C, 0x1D, 0x95)
MAIN = RGBColor(0x6D, 0x28, 0xD9)
GREY = RGBColor(0x55, 0x55, 0x55)
BLACK = RGBColor(0x11, 0x11, 0x11)


def start_a_new_file():
    """Builds the file for the first time, with a short explanation at the top."""
    doc = Document()

    heading = doc.add_paragraph()
    run = heading.add_run("How the agent thinks")
    run.bold = True
    run.font.size = Pt(24)
    run.font.color.rgb = DEEP

    intro = doc.add_paragraph()
    run = intro.add_run(
        "Every time the agent decides which customer complaint should be "
        "handled first, what it did is written down here. The newest run is "
        "always at the bottom.\n\n"
        "Each run walks through the same seven steps, in order, so you can see "
        "where the answer came from rather than just what it was."
    )
    run.font.size = Pt(10.5)
    run.font.color.rgb = GREY

    doc.add_paragraph()
    return doc


def open_or_start():
    """Opens the file if it exists, otherwise starts a new one."""
    if os.path.exists(FILE_PATH):
        return Document(FILE_PATH)
    return start_a_new_file()


def big_line(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(15)
    run.font.color.rgb = DEEP
    return p


def step_heading(doc, number, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(f"Step {number}   {text}")
    run.bold = True
    run.font.size = Pt(11.5)
    run.font.color.rgb = MAIN
    return p


def note(doc, text, indent=0.0, italic=False, small=False):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(indent)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.italic = italic
    run.font.size = Pt(9 if small else 10)
    run.font.color.rgb = GREY if italic or small else BLACK
    return p


def divider(doc):
    p = doc.add_paragraph()
    run = p.add_run("_" * 78)
    run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
    p.paragraph_format.space_after = Pt(10)


def word_for(number):
    return {1: "Low", 2: "Medium", 3: "High", 4: "Critical"}.get(number, str(number))


def write_this_run(result, facts, orders):
    """
    Adds one run to the bottom of the file.

    Everything here is taken from what actually happened. Nothing is
    described that the agent did not really do.
    """
    doc = open_or_start()
    facts_by_id = {f["id"]: f for f in facts}

    divider(doc)

    big_line(doc, "Run on " + datetime.now().strftime("%d %B %Y at %H:%M:%S"))
    note(doc, f"Recorded as {result.get('decision_id', 'not recorded')}",
         italic=True, small=True)

    # ---------- step 1
    step_heading(doc, 1, "Pick up the waiting complaints")
    note(doc, f"{len(facts)} complaints were waiting.")
    for f in facts:
        note(doc, f"{f['id']}   {f['subject']}", indent=0.3, small=True)

    # ---------- step 2
    step_heading(doc, 2, "Look up the three separate systems")
    note(doc, "Each system knows a different part of the story, and none of "
              "them can see what the others see.")
    for f in facts:
        note(doc, f"{f['id']}", indent=0.3, small=True)
        note(doc,
             f"Customer records say: {f['plan']} plan, pays "
             f"\u00a3{f['pays_monthly']:,} a month, has asked "
             f"{f['times_asked']} " + ("time" if f["times_asked"] == 1 else "times"),
             indent=0.6, small=True)
        note(doc,
             f"Monitoring says: {word_for(f['monitoring_says'])}, "
             f"{f['users_hit']} " + ("person" if f["users_hit"] == 1 else "people")
             + " affected, "
             f"work blocked: {'yes' if f['cannot_work'] else 'no'}",
             indent=0.6, small=True)
        note(doc,
             f"The promise clock says: {f['promised_hours']} hours promised, "
             f"{f['hours_left']} left"
             + (", already late" if f["already_late"] else ""),
             indent=0.6, small=True)
        note(doc,
             f"The customer themselves said: {word_for(f['customer_says'])}",
             indent=0.6, small=True)

    # ---------- step 3
    step_heading(doc, 3, "Rank them four different ways")
    note(doc, "Each way is reasonable on its own, and each one is wrong on "
              "its own. They are evidence for the agent to weigh, not "
              "instructions to follow.")

    for name, order in orders.items():
        note(doc, f"By {name}:  " + "  ".join(
            f"{i}. {tid}" for i, tid in enumerate(order, start=1)
        ), indent=0.3, small=True)

    note(doc, "")
    note(doc, "Where the four disagreed most:", indent=0.3)
    for f in facts:
        places = [order.index(f["id"]) + 1 for order in orders.values()]
        gap = max(places) - min(places)
        if gap >= 3:
            note(doc,
                 f"{f['id']} was placed as high as {min(places)} and as low "
                 f"as {max(places)}. This is where judgement was needed.",
                 indent=0.6, small=True)

    # ---------- step 4
    step_heading(doc, 4, "Read what the customer actually wrote")
    note(doc, "The four ways above only look at numbers, so they cannot tell "
              "a password reset apart from a break-in. Both look like one "
              "person with nothing failing. The words are the difference.")
    for f in facts:
        note(doc, f"{f['id']}: \u201c{f['message']}\u201d", indent=0.3, small=True)

    # ---------- step 5
    step_heading(doc, 5, "The AI decides, and says why")
    for position, ticket_id in enumerate(result["order"], start=1):
        f = facts_by_id[ticket_id]
        note(doc, f"{position}. {ticket_id}   {f['subject']}", indent=0.3)
        note(doc, result["reasons"].get(ticket_id, ""), indent=0.6, italic=True)

    note(doc, "")
    note(doc, "The trade-off it made:", indent=0.3)
    note(doc, result["the_trade_off"], indent=0.6, italic=True)

    note(doc, "")
    note(doc, "The closest call:", indent=0.3)
    note(doc, result["hardest_call"], indent=0.6, italic=True)

    ranking = result.get("strategy_ranking", [])
    if ranking:
        note(doc, "")
        note(doc, "How it ranked the four ways of deciding:", indent=0.3)
        for st in ranking:
            note(doc, f"{st['place']}. {st['strategy']}", indent=0.6)
            note(doc, f"why here: {st['why_it_ranks_here']}", indent=0.9, italic=True)
            note(doc, f"gets wrong: {st['what_it_gets_wrong']}", indent=0.9, italic=True)

        note(doc, "")
        note(doc, "Closest call between the four:", indent=0.3)
        note(doc, result.get("why_the_winner_beat_the_runner_up", ""),
             indent=0.6, italic=True)

    # ---------- step 6
    step_heading(doc, 6, "Check the answer before anything happens")
    checks = result["checks"]

    note(doc, "Did every complaint come back, and were the written rules "
              "followed?", indent=0.3)
    if not checks["policy"]:
        note(doc, "Yes. Nothing was dropped and no rule was broken.",
             indent=0.6, italic=True)
    else:
        for c in checks["policy"]:
            note(doc, f"{c['problem']}: {c.get('ticket', '')} "
                      f"{c.get('why', '')}", indent=0.6, italic=True)

    note(doc, "")
    note(doc, "Were the reasons it gave actually true?", indent=0.3)
    wrong = checks.get("reasons_the_data_did_not_support", [])
    if not wrong:
        note(doc, "Yes. Every reason matched the numbers we hold.",
             indent=0.6, italic=True)
    else:
        for c in wrong:
            note(doc, f"{c['ticket']} claimed: {c['the_model_said']}",
                 indent=0.6, small=True)
            note(doc, f"but the data says: {c['but_the_data_says']}",
                 indent=0.6, small=True)

        fixed = checks.get("reasons_we_asked_the_model_to_rewrite", {})
        if fixed:
            note(doc, "")
            note(doc, "So the agent was asked to write those reasons again:",
                 indent=0.3)
            for tid, new_reason in fixed.items():
                note(doc, f"{tid} now reads: {new_reason}",
                     indent=0.6, italic=True)

        if checks.get("needs_a_human_to_look"):
            note(doc, "Some reasons still did not match. A person should look "
                      "at this before it is acted on.", indent=0.6, italic=True)

    note(doc, "")
    note(doc, "How solid was the answer? One number was changed at a time and "
              "everything ranked again, to see which places were close calls.",
         indent=0.3)
    for r in result.get("how_solid_the_scoring_is", {}).values():
        line = (f"{r['id']}  {r['how_solid']}  "
                f"({r['small_changes_that_move_it']} of "
                f"{r['small_changes_tried']} small changes move it)")
        note(doc, line, indent=0.6, small=True)
        for m in r["what_would_move_it"][:2]:
            note(doc, f"if {m['if']}, it moves from {m['moves_from']} "
                      f"to {m['moves_to']}", indent=0.9, small=True, italic=True)

    # ---------- what happens next
    step_heading(doc, 7, "Hand them to a person")
    for position, ticket_id in enumerate(result["order"], start=1):
        when = "now" if position == 1 else ("next" if position <= 3 else "queued")
        solid = result.get("how_solid_the_scoring_is", {}).get(ticket_id, {})
        tail = ("   (the numbers were not settled about this one)"
                if solid.get("worth_a_second_look") else "")
        note(doc, f"{ticket_id} goes to a person {when}{tail}",
             indent=0.3, small=True)

    doc.add_paragraph()
    doc.save(FILE_PATH)

    return FILE_PATH
# ----- THE SAME STORY, AS A MARKDOWN FILE -----
#
# GitHub cannot show a Word file in the browser, it only downloads it.
# So we write the same thing as markdown too, which GitHub displays
# straight away. Same content, two formats, for two kinds of reader.

def write_this_run_as_markdown(result, facts, orders):
    """Adds one run to the bottom of the markdown file."""
    facts_by_id = {f["id"]: f for f in facts}
    lines = []

    if not os.path.exists(MARKDOWN_PATH):
        lines.append("# How the agent thinks\n")
        lines.append(
            "Every time the agent decides which customer complaint should be "
            "handled first, what it did is written down here. The newest run is "
            "always at the bottom.\n"
        )
        lines.append(
            "Each run walks through the same seven steps, in order, so you can "
            "see where the answer came from rather than just what it was.\n"
        )

    lines.append("\n---\n")
    lines.append(f"## Run on {datetime.now().strftime('%d %B %Y at %H:%M:%S')}\n")
    lines.append(f"*Recorded as {result.get('decision_id', 'not recorded')}*\n")

    # ---------- 1
    lines.append("\n### Step 1  Pick up the waiting complaints\n")
    lines.append(f"{len(facts)} complaints were waiting.\n")
    for f in facts:
        lines.append(f"- `{f['id']}`  {f['subject']}")

    # ---------- 2
    lines.append("\n### Step 2  Look up the three separate systems\n")
    lines.append("Each system knows a different part of the story, and none of "
                 "them can see what the others see.\n")
    lines.append("| Ticket | Customer records | Monitoring | The promise clock | They said |")
    lines.append("|---|---|---|---|---|")
    for f in facts:
        blocked = "blocked" if f["cannot_work"] else "not blocked"
        late = ", already late" if f["already_late"] else ""
        lines.append(
            f"| `{f['id']}` "
            f"| {f['plan']}, £{f['pays_monthly']:,}/mo, asked {f['times_asked']}x "
            f"| {word_for(f['monitoring_says'])}, {f['users_hit']} affected, {blocked} "
            f"| {f['promised_hours']}h promised, {f['hours_left']}h left{late} "
            f"| {word_for(f['customer_says'])} |"
        )

    # ---------- 3
    lines.append("\n### Step 3  Rank them four different ways\n")
    lines.append("Each way is reasonable on its own, and each one is wrong on its "
                 "own. They are evidence for the agent to weigh, not instructions "
                 "to follow.\n")
    for name, order in orders.items():
        lines.append(f"- **{name}**: " + " → ".join(order))

    lines.append("\nWhere the four disagreed most:\n")
    for f in facts:
        places = [order.index(f["id"]) + 1 for order in orders.values()]
        if max(places) - min(places) >= 3:
            lines.append(f"- `{f['id']}` was placed as high as {min(places)} and as "
                         f"low as {max(places)}. This is where judgement was needed.")

    # ---------- 4
    lines.append("\n### Step 4  Read what the customer actually wrote\n")
    lines.append("The four ways above only look at numbers, so they cannot tell a "
                 "password reset apart from a break-in. Both look like one person "
                 "with nothing failing. The words are the difference.\n")
    for f in facts:
        lines.append(f"- `{f['id']}`: \"{f['message']}\"")

    # ---------- 5
    lines.append("\n### Step 5  The AI decides, and says why\n")
    for position, ticket_id in enumerate(result["order"], start=1):
        f = facts_by_id[ticket_id]
        lines.append(f"**{position}. `{ticket_id}`  {f['subject']}**  ")
        lines.append(f"{result['reasons'].get(ticket_id, '')}\n")

    conflicts = result.get("conflicts_i_noticed", [])
    if conflicts:
        lines.append("\n**The contradictions it spotted:**\n")
        for c in conflicts:
            lines.append(f"- `{', '.join(c.get('tickets', []))}` "
                         f"{c.get('what_disagreed', '')}. "
                         f"It believed {c.get('which_i_believed', '')}, "
                         f"because {c.get('why', '')}")

    lines.append(f"\n**The trade-off it made:** {result.get('the_trade_off', '')}\n")
    lines.append(f"**The closest call:** {result.get('hardest_call', '')}\n")

    ranking = result.get("strategy_ranking", [])
    if ranking:
        lines.append("\n**How it ranked the four ways of deciding:**\n")
        lines.append("| Place | Way | Why here | What it gets wrong |")
        lines.append("|---|---|---|---|")
        for st in ranking:
            lines.append(f"| {st.get('place', '')} | {st.get('strategy', '')} "
                         f"| {st.get('why_it_ranks_here', '')} "
                         f"| {st.get('what_it_gets_wrong', '')} |")

    # ---------- 6
    lines.append("\n### Step 6  Check the answer before anything happens\n")
    checks = result["checks"]

    lines.append("**Did every complaint come back, and were the written rules "
                 "followed?**  ")
    if not checks["policy"]:
        lines.append("Yes. Nothing was dropped and no rule was broken.\n")
    else:
        for c in checks["policy"]:
            lines.append(f"{c['problem']}: {c.get('ticket', '')} {c.get('why', '')}\n")

    lines.append("**Were the reasons it gave actually true?**  ")
    wrong = checks.get("reasons_the_data_did_not_support", [])
    if not wrong:
        lines.append("Yes. Every reason matched the numbers we hold.\n")
    else:
        for c in wrong:
            lines.append(f"- `{c['ticket']}` claimed: *{c['the_model_said']}*")
            lines.append(f"  but the data says: **{c['but_the_data_says']}**")

        fixed = checks.get("reasons_we_asked_the_model_to_rewrite", {})
        if fixed:
            lines.append("\nSo the agent was asked to write those reasons again:\n")
            for tid, new_reason in fixed.items():
                lines.append(f"- `{tid}` now reads: {new_reason}")

        if checks.get("needs_a_human_to_look"):
            lines.append("\nSome reasons still did not match. A person should look "
                         "at this before it is acted on.\n")

    lines.append("\n**How solid was the answer?** One number was changed at a time "
                 "and everything ranked again, to see which places were close calls.\n")
    lines.append("| Ticket | How solid | Changes that move it |")
    lines.append("|---|---|---|")
    for r in result.get("how_solid_the_scoring_is", {}).values():
        lines.append(f"| `{r['id']}` | {r['how_solid']} | "
                     f"{r['small_changes_that_move_it']} of "
                     f"{r['small_changes_tried']} |")

    # ---------- 7
    lines.append("\n### Step 7  Hand them to a person\n")
    for position, ticket_id in enumerate(result["order"], start=1):
        when = "now" if position == 1 else ("next" if position <= 3 else "queued")
        solid = result.get("how_solid_the_scoring_is", {}).get(ticket_id, {})
        tail = ("  *(the numbers were not settled about this one)*"
                if solid.get("worth_a_second_look") else "")
        lines.append(f"- `{ticket_id}` goes to a person **{when}**{tail}")

    lines.append("")

    with open(MARKDOWN_PATH, "a", encoding="utf-8") as f:
        f.write("\n".join(lines) + "\n")

    return MARKDOWN_PATH
